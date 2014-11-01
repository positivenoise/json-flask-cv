import os, time
from flask import Flask, render_template, json, make_response, g
from flask_weasyprint import HTML, render_pdf
from docx import Document

os.environ['TZ'] = 'AEST-10'
time.tzset()

def import_json():
    f =  open("resume.json")
    jdat = f.read()
    json_data = json.loads(jdat)
    f.close()
    return json_data

def generate_docx(json_data):

    document = Document("static/template.docx")
    document.add_heading(json_data["basics"]["name"], 0)
    document.add_paragraph(json_data["basics"]["email"] + ' | ' + json_data["basics"]["location"] + ' | ' + json_data["basics"]["phone"] + ' | ' + 'Generated on: ' + time.strftime("%a, %d %b %Y %H:%M:%S AEST"))
    document.add_heading('Education', level=1)

    for s in json_data["education"]:
        document.add_paragraph(s["name"] + " - Obtained: " + s["date"] , style='ListBullet')
    document.add_heading('Skills', level=1)

    for s in json_data["skills"]:
        document.add_heading(s["type"], level=2)
        for dat in s["keywords"]:
            document.add_paragraph(dat, style='ListBullet')
    document.add_heading("Experience", level=2)

    for s in json_data["work"]:
        document.add_paragraph(s["company"] + " | " + s["position"] + " | from " + s["startDate"] + " to " + s["endDate"])
        document.add_paragraph(s["summary"])
        document.add_heading("Highlights", level=2)
        for x in s["highlights"]:
            document.add_paragraph(x, style='ListBullet')

    document.add_page_break()
    document.save("static/resume.docx")
    return document

def json2xml(json_obj, line_padding=""):
    rl = list()
    json_obj_type = type(json_obj)

    if json_obj_type is list:
        for sub_elem in json_obj:
            rl.append(json2xml(sub_elem, line_padding))

        return "\n".join(rl)

    if json_obj_type is dict:
        for tag in json_obj:
            sub_obj = json_obj[tag]
            rl.append("%s<%s>" % (line_padding, tag))
            rl.append(json2xml(sub_obj, "\t" + line_padding))
            rl.append("%s</%s>" % (line_padding, tag))
        return "\n".join(rl)

    return "%s%s" % (line_padding, json_obj)

app = Flask(__name__)

@app.before_request
def before_request():
    g.json_data = import_json()
@app.route('/')
def index():
    return render_template('index.html', json = g.json_data, generated = time.strftime("%a, %d %b %Y %H:%M:%S AEST"))

@app.route('/xml')
def xml():
    return render_template('xml.html', xml = json2xml(g.json_data), obj = g.json_data)

@app.route('/json')
def jsonpage():
    return render_template('json.html', json = json.dumps(g.json_data, indent=4, separators=(',', ': ')), obj = g.json_data)

@app.route('/pdf')
def pdf():
    html = render_template('pdf.html', json=g.json_data, generated = time.strftime("%a, %d %b %Y %H:%M:%S AEST"))
    return render_pdf(HTML(string=html))

@app.route('/download/json')
def download_json():
    response = make_response(json.dumps(g.json_data, indent=4, separators=(',', ': ')))
    response.headers["Content-Disposition"] = "attachment; filename=resume.json"
    return response

@app.route('/download/xml')
def download_xml():
    response = make_response(json2xml(g.json_data))
    response.headers["Content-Disposition"] = "attachment; filename=resume.xml"
    return response

@app.route('/docx')
def download_docx():
    generate_docx(g.json_data)
    with open("static/resume.docx", 'r') as f:
        body = f.read()
    response = make_response(body)
    response.headers["Content-Disposition"] = "attachment; filename=resume.docx"
    return response

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
